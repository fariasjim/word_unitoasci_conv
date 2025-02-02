from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from tkinter import Tk
from tkinter.filedialog import askopenfilename, asksaveasfilename
import converter  # Import the converter module
import re
global replaced_runs
replaced_runs=[]

def contains_unicode(text):
    # Check if the text contains Unicode characters
    return bool(re.search(r'[^\x00-\x7F]', text))

def replace_and_highlight(doc_path, save_path):
    doc = Document(doc_path)
    
    # Create an instance of the Unicode class from converter.py
    unicode_converter = converter.Unicode()
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if contains_unicode(run.text):
                if contains_unicode(run.text) and run.font.name == "SolaimanLipi":
            # Convert the text from Unicode to Bijoy
                    converted_text = unicode_converter.convertUnicodeToBijoy(run.text)
                    if run.text != converted_text:
                        run.text = converted_text
                        run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE  # BLUE highlight
                        replaced_runs.append(run)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if contains_unicode(run.text):
                            if contains_unicode(run.text) and run.font.name == "SolaimanLipi":
                        # Convert the text from Unicode to Bijoy
                                converted_text = unicode_converter.convertUnicodeToBijoy(run.text)
                                if run.text != converted_text:
                                    run.text = converted_text
                                    run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE  # BLUE highlight
                                    replaced_runs.append(run)

    for run in replaced_runs:
        run.font.name = "SutonnyMJ"
    
    # Save the document
    doc.save(save_path)

# Create a Tkinter root window
root = Tk()
root.withdraw()  # Hide the root window

# Open file dialog to select the input document
doc_path = askopenfilename(title="Select the Input Document", filetypes=[("Word Documents", "*.docx")])
if not doc_path:
    import tkinter as tk
    from tkinter import messagebox

    def show_info_dialog(custom_text):
        # Create a Tkinter root window
        root = tk.Tk()
        root.withdraw()  # Hide the root window

        # Show an information dialog box with custom text
        messagebox.showinfo("Information", custom_text)

    # Custom text to display in the dialog box
    custom_text = "No Document Selected. Exiting"

    # Show the information dialog box
    show_info_dialog(custom_text)
    exit()

# Open file dialog to select the save location
save_path = asksaveasfilename(title="Save the Output Document", defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
if not save_path:
    import tkinter as tk
    from tkinter import messagebox

    def show_info_dialog(custom_text):
        # Create a Tkinter root window
        root = tk.Tk()
        root.withdraw()  # Hide the root window

        # Show an information dialog box with custom text
        messagebox.showinfo("Information", custom_text)

    # Custom text to display in the dialog box
    custom_text = "Can't save file. Exiting."

    # Show the information dialog box
    show_info_dialog(custom_text)
    exit()

replace_and_highlight(doc_path, save_path)

import tkinter as tk
from tkinter import messagebox

def show_info_dialog(custom_text):
    # Create a Tkinter root window
    root = tk.Tk()
    root.withdraw()  # Hide the root window

    # Show an information dialog box with custom text
    messagebox.showinfo("Information", custom_text)

# Custom text to display in the dialog box
custom_text = "Document converted successfully. This script was been compiled by Farias Hamid Jim"

# Show the information dialog box
show_info_dialog(custom_text)

print(replaced_runs)
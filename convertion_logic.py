from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import converter  # Import the converter module
import re
from tkinter import messagebox

runs = []

def contains_unicode(text):
    # Check if the text contains Unicode characters
    return bool(re.search(r'[^\x00-\x7F]', text))

def replace_and_highlight(doc_path, save_path):
    doc = Document(doc_path)
    
    # Create an instance of the Unicode class from converter.py
    unicode_converter = converter.Unicode()
    
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if contains_unicode(run.text):
                runs.append(run)
                if contains_unicode(run.text) and run.font.name == "SolaimanLipi":
            # Convert the text from Unicode to Bijoy
                    converted_text = unicode_converter.convertUnicodeToBijoy(run.text)
                    if run.text != converted_text:
                        run.text = converted_text
                        run.font.highlight_color = WD_COLOR_INDEX.TEAL  # Teal highlight
                        run.font.name = "SutonnyMJ"  # Set font to Bijoy
                        runs.append(run)
                        

    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if contains_unicode(run.text):
                            if contains_unicode(run.text) and run.font.name == "SolaimanLipi":
                        # Convert the text from Unicode to Bijoy
                                converted_text = unicode_converter.convertUnicodeToBijoy(run.text)
                                if run.text != converted_text:
                                    run.text = converted_text
                                    run.font.highlight_color = WD_COLOR_INDEX.TEAL  # Teal highlight
                                    run.font.name = "SutonnyMJ" # Set font to Bijoy
                                    runs.append(run)

    for run in runs:
        try:
            run.font.name = "SutonnyMJ"
        except Exception as e:
            print(f"Error setting font: {e}")
    # Save the document
    doc.save(save_path)
    messagebox.showinfo(f"Total words converted: {len(runs)}", "Conversion Complete")
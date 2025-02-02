import sys
import os
from PyQt5.QtWidgets import QApplication, QMainWindow, QVBoxLayout, QWidget, QPushButton, QFileDialog, QTextEdit
from PyQt5.QtGui import QFont, QFontDatabase, QTextCursor
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Pt
from docx2txt import process
from PIL import Image

class WordViewer(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Word Document Viewer")
        self.setGeometry(100, 100, 900, 600)

        # QTextEdit for displaying formatted text
        self.text_edit = QTextEdit(self)
        self.text_edit.setReadOnly(False)

        # Buttons
        self.open_button = QPushButton('Open Word Document', self)
        self.open_button.clicked.connect(self.open_file)

        self.save_button = QPushButton('Save as Word Document', self)
        self.save_button.clicked.connect(self.save_file)

        # Layout
        layout = QVBoxLayout()
        layout.addWidget(self.open_button)
        layout.addWidget(self.save_button)
        layout.addWidget(self.text_edit)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

        self.image_paths = {}

    def open_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Open Word Document", "", "Word Documents (*.docx)")
        if file_path:
            self.load_word_document(file_path)

    def load_word_document(self, file_path):
        doc = Document(file_path)
        text = ""

        # Temporary folder for images
        img_folder = "extracted_images"
        os.makedirs(img_folder, exist_ok=True)

        for element in doc.element.body:
            if element.tag.endswith("p"):  # Paragraph
                para = next(p for p in doc.paragraphs if p._element is element)
                para_text = ""
                para_style = ""

                # Handle paragraph alignment
                if para.alignment == 1:
                    para_style += "text-align: center; "
                elif para.alignment == 2:
                    para_style += "text-align: right; "
                elif para.alignment == 3:
                    para_style += "text-align: justify; "

                for run in para.runs:
                    run_style = ""

                    # Apply font name (if available)
                    if run.font.name:
                        font_id = QFontDatabase.addApplicationFont(run.font.name)
                        font_family = QFontDatabase.applicationFontFamilies(font_id)
                        if font_family:
                            run_style += f"font-family: '{font_family[0]}'; "

                    # Check font properties
                    if run.font.size:
                        run_style += f"font-size: {int(run.font.size.pt)}pt; "
                    if run.font.color and run.font.color.rgb:
                        color_hex = run.font.color.rgb
                        run_style += f"color: #{color_hex}; "
                    if run.bold:
                        run_style += "font-weight: bold; "
                    if run.italic:
                        run_style += "font-style: italic; "
                    if run.underline:
                        run_style += "text-decoration: underline; "

                    para_text += f"<span style='{run_style}'>{run.text}</span>"

                text += f"<div style='{para_style}'>{para_text}</div>"

            elif element.tag.endswith("tbl"):  # Table
                table = next(t for t in doc.tables if t._element is element)
                text += "<table border='1' cellspacing='0' cellpadding='5' style='border-collapse: collapse; width: 100%;'>"
                for row in table.rows:
                    text += "<tr>"
                    for cell in row.cells:
                        text += f"<td style='padding: 5px;'>{cell.text}</td>"
                    text += "</tr>"
                text += "</table>"

        # Extract images
        extracted_images = process(file_path, img_folder)

        # Insert images into the text
        for img_file in os.listdir(img_folder):
            img_path = os.path.join(img_folder, img_file)
            img_tag = f'<img src="{img_path}" width="300">'
            text += f"<br>{img_tag}<br>"
            self.image_paths[img_file] = img_path

        # Set HTML content in QTextEdit
        self.text_edit.setHtml(text)

    def save_file(self):
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Word Document", "", "Word Documents (*.docx)")
        if save_path:
            self.save_to_word(save_path)

    def save_to_word(self, save_path):
        doc = Document()

        # Convert HTML back to docx
        cursor = QTextCursor(self.text_edit.document())
        while not cursor.atEnd():
            cursor.movePosition(QTextCursor.NextBlock)
            block = cursor.block()
            text = block.text()

            if text.strip().startswith("<img"):  # If it's an image
                img_path = self.image_paths.get(text.strip().split('"')[1], None)
                if img_path:
                    doc.add_picture(img_path, width=Pt(200))
            else:
                para = doc.add_paragraph(text)
                para.alignment = Qt.AlignLeft  # Apply basic left alignment

        doc.save(save_path)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    viewer = WordViewer()
    viewer.show()
    sys.exit(app.exec_())
